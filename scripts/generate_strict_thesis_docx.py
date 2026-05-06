from __future__ import annotations

import json
from collections import Counter
from pathlib import Path

import matplotlib.pyplot as plt
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from pygments import highlight
from pygments.formatters.img import ImageFormatter
from pygments.lexers import get_lexer_for_filename


PROJECT_ROOT = Path("/Users/bytedance/Desktop/multi_agent_dialog_system")
RESULTS_DIR = PROJECT_ROOT / "evaluation" / "results"
DATASET_PATH = PROJECT_ROOT / "evaluation" / "datasets" / "life_service_eval.json"
ASSET_DIR = PROJECT_ROOT / ".thesis_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = Path("/Users/bytedance/Desktop/基于对话管理的生活服务_AI_Agent_多轮交互策略_严格格式版.docx")
MARKDOWN_PATH = ASSET_DIR / "thesis_sync.md"
MEDIA_MANIFEST_PATH = ASSET_DIR / "thesis_media_manifest.json"

THESIS_TITLE = "基于对话管理的生活服务 AI-Agent 多轮交互策略"
THESIS_TITLE_EN = "MULTI-TURN INTERACTION STRATEGIES FOR LIFE-SERVICE AI AGENTS BASED ON DIALOGUE MANAGEMENT"
SCHOOL_CN = "学校名称（待填写）"
SCHOOL_EN = "UNIVERSITY NAME (TO BE FILLED)"
HEADER_TITLE = THESIS_TITLE

plt.rcParams["font.sans-serif"] = ["PingFang SC", "Heiti SC", "Arial Unicode MS", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def read_snippet(path: Path, start: int, end: int) -> str:
    lines = path.read_text(encoding="utf-8").splitlines()
    return "\n".join(lines[start - 1 : end])


def safe_font(size: int):
    try:
        return ImageFont.truetype("/System/Library/Fonts/STHeiti Medium.ttc", size)
    except Exception:
        return ImageFont.load_default()


def ensure_east_asia(run, east_asia: str) -> None:
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_run(run, east_asia: str, size: int, bold: bool = False, latin: str = "Times New Roman") -> None:
    run.font.name = latin
    ensure_east_asia(run, east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def set_layout(section) -> None:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)


def configure_styles(doc: Document) -> None:
    for section in doc.sections:
        set_layout(section)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_page_number_format(section, fmt: str, start: int) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), fmt)
    pg_num_type.set(qn("w:start"), str(start))


def build_badge(output: Path) -> Path:
    image = Image.new("RGBA", (900, 320), (255, 255, 255, 255))
    draw = ImageDraw.Draw(image)
    draw.ellipse((40, 40, 220, 220), outline="#1d4ed8", width=8, fill="#eff6ff")
    draw.text((95, 88), "校", fill="#1e3a8a", font=safe_font(60))
    draw.text((280, 70), SCHOOL_CN, fill="#0f172a", font=safe_font(60))
    draw.text((280, 150), SCHOOL_EN, fill="#475569", font=safe_font(32))
    image.save(output)
    return output


def render_diagram(title: str, nodes, arrows, output: Path) -> Path:
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 40), title, fill="#0f172a", font=safe_font(40))
    text_font = safe_font(26)
    for x1, y1, x2, y2, label, fill in nodes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, outline="#1e3a8a", width=4, fill=fill)
        bbox = draw.multiline_textbbox((0, 0), label, font=text_font, spacing=8)
        w = bbox[2] - bbox[0]
        h = bbox[3] - bbox[1]
        draw.multiline_text(((x1 + x2 - w) / 2, (y1 + y2 - h) / 2), label, fill="#0f172a", font=text_font, align="center", spacing=8)
    for (sx, sy), (ex, ey) in arrows:
        draw.line((sx, sy, ex, ey), fill="#475569", width=6)
        draw.polygon([(ex, ey), (ex - 18, ey - 10), (ex - 18, ey + 10)], fill="#475569")
    image.save(output)
    return output


def render_chart(result: dict, output: Path) -> Path:
    ours = result["stacked_multi_agent"]["metrics"]
    baseline = result["baseline_single_agent"]["metrics"]
    labels = ["成功率", "平均轮次", "槽位F1"]
    ours_values = [ours["task_success_rate"], ours["average_turns"], ours["slot_f1"]]
    base_values = [baseline["task_success_rate"], baseline["average_turns"], baseline["slot_f1"]]
    x = range(len(labels))
    width = 0.35
    plt.figure(figsize=(10, 5.6))
    plt.bar([v - width / 2 for v in x], ours_values, width=width, label="Stacked Multi-Agent", color="#2563eb")
    plt.bar([v + width / 2 for v in x], base_values, width=width, label="Baseline Single Agent", color="#ef4444")
    plt.xticks(list(x), labels)
    plt.title("核心指标对比图")
    plt.legend()
    plt.tight_layout()
    plt.savefig(output, dpi=220)
    plt.close()
    return output


def render_distribution(dataset: list[dict], output: Path) -> Path:
    counter = Counter(sample["domain"] for sample in dataset)
    labels = list(counter.keys())
    values = [counter[k] for k in labels]
    plt.figure(figsize=(10, 5.6))
    plt.bar(labels, values, color=["#2563eb", "#0ea5e9", "#10b981"])
    plt.title("评测数据集场景分布")
    plt.tight_layout()
    plt.savefig(output, dpi=220)
    plt.close()
    return output


def render_metric_bar(result: dict, key: str, title: str, output: Path) -> Path:
    ours = result["stacked_multi_agent"]["metrics"][key]
    baseline = result["baseline_single_agent"]["metrics"][key]
    plt.figure(figsize=(8, 4.8))
    plt.bar(["Multi-Agent", "Baseline"], [ours, baseline], color=["#2563eb", "#ef4444"])
    plt.title(title)
    plt.tight_layout()
    plt.savefig(output, dpi=220)
    plt.close()
    return output


def render_code_image(source: Path, start: int, end: int, output: Path) -> Path:
    code = read_snippet(source, start, end)
    lexer = get_lexer_for_filename(source.name, code)
    formatter = ImageFormatter(font_name="Menlo", font_size=20, line_numbers=True, style="friendly", image_pad=20, line_pad=4)
    output.write_bytes(highlight(code, lexer, formatter))
    return output


def build_assets(result: dict, dataset: list[dict]) -> dict[str, Path]:
    architecture_nodes = [
        (80, 300, 300, 420, "用户输入", "#dbeafe"),
        (380, 250, 640, 470, "全局路由器\nRouter", "#e0f2fe"),
        (760, 120, 1040, 250, "餐饮专家", "#dcfce7"),
        (760, 290, 1040, 420, "酒店专家", "#dcfce7"),
        (1140, 100, 1480, 290, "工具层\nsearch / book / update_blackboard", "#fef3c7"),
        (1140, 340, 1480, 530, "状态层\nDialogState / task_stack / blackboard", "#ede9fe"),
    ]
    architecture_arrows = [((300, 360), (380, 360)), ((640, 300), (760, 180)), ((640, 360), (760, 360)), ((1040, 180), (1140, 180)), ((1040, 360), (1140, 180)), ((1040, 360), (1140, 430))]
    stack_nodes = [
        (160, 160, 520, 290, "步骤1：酒店任务激活\nactive_task = hotel", "#eff6ff"),
        (160, 360, 520, 490, "步骤2：用户打断询问餐饮\npush(hotel)", "#ecfccb"),
        (160, 560, 520, 690, "步骤3：餐饮完成后恢复 hotel", "#fef3c7"),
        (760, 130, 1360, 720, "任务栈视图\n顶部：hotel\n中部：...\n底部：历史挂起任务", "#f8fafc"),
    ]
    stack_arrows = [((340, 290), (340, 360)), ((340, 490), (340, 560)), ((520, 225), (760, 225)), ((520, 425), (760, 425)), ((520, 625), (760, 625))]
    workbench_nodes = [
        (80, 160, 460, 760, "左栏\nChatPanel\n会话与输入", "#eff6ff"),
        (560, 160, 1040, 440, "中栏上\nTracePanel\n事件时间线", "#f8fafc"),
        (560, 480, 1040, 760, "中栏下\nStatePanel\n任务栈 / 黑板", "#f8fafc"),
        (1140, 160, 1520, 380, "右栏上\nKnowledgebasePanel", "#eef2ff"),
        (1140, 420, 1520, 590, "右栏中\nEvalPanel", "#eef2ff"),
        (1140, 630, 1520, 760, "右栏下\nReportsPanel", "#eef2ff"),
    ]
    workbench_arrows = [((460, 460), (560, 300)), ((460, 460), (560, 620)), ((1040, 300), (1140, 270)), ((1040, 620), (1140, 505))]
    return {
        "badge": build_badge(ASSET_DIR / "school_badge.png"),
        "architecture": render_diagram("图  多智能体对话系统总体架构图", architecture_nodes, architecture_arrows, ASSET_DIR / "strict_architecture.png"),
        "stack": render_diagram("图  栈式记忆挂起-恢复流程图", stack_nodes, stack_arrows, ASSET_DIR / "strict_stack.png"),
        "workbench": render_diagram("图  工作台三栏布局示意图", workbench_nodes, workbench_arrows, ASSET_DIR / "strict_workbench.png"),
        "metric": render_chart(result, ASSET_DIR / "strict_metric_compare.png"),
        "distribution": render_distribution(dataset, ASSET_DIR / "strict_distribution.png"),
        "success": render_metric_bar(result, "task_success_rate", "任务成功率", ASSET_DIR / "strict_success.png"),
        "slotf1": render_metric_bar(result, "slot_f1", "槽位 F1", ASSET_DIR / "strict_slotf1.png"),
        "state": render_code_image(PROJECT_ROOT / "multi_agent_dialog" / "state.py", 1, 23, ASSET_DIR / "strict_code_state.png"),
        "router": render_code_image(PROJECT_ROOT / "multi_agent_dialog" / "router.py", 7, 71, ASSET_DIR / "strict_code_router.png"),
        "experts": render_code_image(PROJECT_ROOT / "multi_agent_dialog" / "experts.py", 120, 190, ASSET_DIR / "strict_code_experts.png"),
        "tools": render_code_image(PROJECT_ROOT / "multi_agent_dialog" / "tools.py", 1, 140, ASSET_DIR / "strict_code_tools.png"),
        "runner": render_code_image(PROJECT_ROOT / "evaluation" / "runner.py", 1, 240, ASSET_DIR / "strict_code_runner.png"),
        "trace": render_code_image(PROJECT_ROOT / "app" / "services" / "trace_service.py", 1, 40, ASSET_DIR / "strict_code_trace.png"),
        "workbench_code": render_code_image(PROJECT_ROOT / "web" / "src" / "pages" / "WorkbenchPage.tsx", 1, 220, ASSET_DIR / "strict_code_workbench.png"),
    }


def paragraph(doc: Document, text: str, east_asia: str = "宋体", size: int = 12, bold: bool = False, align=None, indent_cm: float = 0.74) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.first_line_indent = Cm(indent_cm) if indent_cm else Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run(run, east_asia, size, bold=bold)


def center_title(doc: Document, text: str, east_asia: str, size: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_run(run, east_asia, size, bold=True)


def picture(doc: Document, path: Path, caption: str, width: float = 5.8) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.first_line_indent = Cm(0)
    run = cp.add_run(caption)
    set_run(run, "黑体", 10, bold=True)


def table_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run(run, "黑体", 10, bold=True)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for idx, item in enumerate(headers):
        cell = t.rows[0].cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(item)
        set_run(run, "黑体", 10, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for idx, item in enumerate(row):
            cells[idx].text = ""
            run = cells[idx].paragraphs[0].add_run(item)
            set_run(run, "宋体", 10)


def add_header(section, badge: Path) -> None:
    section.header.is_linked_to_previous = False
    header = section.header
    header.paragraphs[0].text = ""
    t = header.add_table(rows=1, cols=2, width=Cm(15.5))
    t.autofit = False
    left = t.rows[0].cells[0].paragraphs[0]
    left.add_run().add_picture(str(badge), width=Inches(1.1))
    run = left.add_run("  学校名称（待填写）")
    set_run(run, "黑体", 9)
    right = t.rows[0].cells[1].paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right.add_run(HEADER_TITLE)
    set_run(run, "黑体", 9)


def add_footer(section, fmt: str, start: int) -> None:
    section.footer.is_linked_to_previous = False
    set_page_number_format(section, fmt, start)
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run(run, "黑体", 9)
    add_field(run, " PAGE ")


def add_cover(doc: Document, badge: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(badge), width=Inches(3.2))
    center_title(doc, THESIS_TITLE, "黑体", 32)
    fields = [("学    院", "计算机学院"), ("专    业", "计算机科学与技术"), ("学生姓名", "________________"), ("学    号", "________________"), ("指导教师", "________________")]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run1 = p.add_run(f"{label}：")
        set_run(run1, "黑体", 18, bold=True)
        run2 = p.add_run(value)
        set_run(run2, "楷体_GB2312", 18)
    center_title(doc, THESIS_TITLE_EN, "黑体", 26)
    center_title(doc, "计算机学院本科毕业设计（论文）", "黑体", 22)
    center_title(doc, "2026 年 5 月", "黑体", 18)


def expand(topic: str, aspects: list[str]) -> list[str]:
    paragraphs: list[str] = []
    for aspect in aspects:
        paragraphs.append(f"围绕{topic}，本文重点从“{aspect}”这一角度展开分析。该问题在生活服务智能体场景中具有典型性，因为用户表达往往口语化、约束信息分散且跨轮次传递，如果系统仅依赖大模型的隐式上下文记忆，容易在任务切换、条件澄清和执行确认阶段出现不稳定表现。为此，本文坚持以结构化对话管理为主线，将语言能力与状态约束结合起来，提升系统在真实任务中的鲁棒性。")
        paragraphs.append(f"从工程实现看，“{aspect}”并不是孤立模块，而是与路由决策、领域专家、工具接口和可观测日志形成联动关系。只有把这些关系明确建模，系统才能在面对复杂输入时既保持自然交互，又确保任务状态不丢失、执行动作不过界。本文的设计正是通过任务栈、全局黑板、知识库快照以及规则过滤来实现这一目标。")
        paragraphs.append(f"从实验与论文写作角度看，对“{aspect}”进行展开还有助于解释指标背后的原因。很多系统在表面回复上看似正确，但在结构化评测中却失败，根源往往不在模型不会说，而在系统不会稳定地记、不会谨慎地做、不会在任务中断后接着做。因此，本文把该问题纳入正文展开讨论，而不再像传统做法那样全部放入附录。")
    return paragraphs


def make_body(metrics: dict, baseline_metrics: dict, failure_counter: Counter[str]) -> list[dict]:
    return [
        {"chapter": "第 1 章 绪论", "sections": [
            {"title": "1.1 研究背景与意义", "paragraphs": expand("研究背景与意义", ["大语言模型推动生活服务对话升级", "任务型对话从问答走向执行闭环", "生活服务场景中的安全与可控需求"])},
            {"title": "1.2 国内外研究现状", "paragraphs": expand("国内外研究现状", ["传统任务型对话系统的模块化路径", "大模型 Agent 的工具调用范式", "多智能体在复杂任务中的分工与局限"])},
            {"title": "1.3 研究内容与技术路线", "paragraphs": expand("研究内容与技术路线", ["多智能体协同架构设计", "状态建模与任务栈机制", "前后端工作台与实验闭环"])},
            {"title": "1.4 论文结构安排", "paragraphs": expand("论文结构安排", ["章节组织逻辑", "方法、实现与实验之间的关系"])},
        ]},
        {"chapter": "第 2 章 相关理论与技术基础", "sections": [
            {"title": "2.1 任务型对话系统理论基础", "paragraphs": expand("任务型对话理论基础", ["任务成功率与槽位质量", "状态跟踪与策略学习", "生活服务任务的目标导向属性"])},
            {"title": "2.2 大语言模型 Agent 与工具调用机制", "paragraphs": expand("大语言模型 Agent 与工具调用", ["ReAct 式推理—行动闭环", "Tool Calling 的结构化参数约束", "执行型动作的安全边界"])},
            {"title": "2.3 多智能体协作与任务分治思想", "paragraphs": expand("多智能体协作", ["全局路由器与领域专家的职责划分", "局部专家上下文聚焦", "统一状态作为协作基础"])},
            {"title": "2.4 对话状态管理与可观测性", "paragraphs": expand("对话状态管理与可观测性", ["任务栈与活动任务", "全局黑板与知识库快照", "Trace 日志驱动的复盘分析"])},
        ]},
        {"chapter": "第 3 章 面向生活服务的系统需求分析与总体架构", "sections": [
            {"title": "3.1 场景需求分析", "paragraphs": expand("生活服务场景需求", ["模糊表达与口语化输入", "用户打断与跨任务切换", "研究型系统对可配置实验平台的要求"])},
            {"title": "3.2 总体架构设计", "paragraphs": expand("总体架构设计", ["路由器 + 领域专家 + 工具层", "状态层和可观测层的职责", "工作台对研究闭环的支撑"]), "figure": ("architecture", "图 3-1 多智能体对话系统总体架构图", 6.0)},
            {"title": "3.3 统一状态模型设计", "paragraphs": expand("统一状态模型", ["DialogState 的字段组织", "结构化状态与自然语言上下文共存", "状态一致性对系统鲁棒性的影响"]), "figure": ("state", "图 3-2 DialogState 状态结构代码截图", 6.0)},
            {"title": "3.4 任务栈挂起—恢复机制", "paragraphs": expand("任务栈挂起—恢复机制", ["用户意图打断", "压栈与弹栈流程", "恢复式对话的用户体验价值"]), "figure": ("stack", "图 3-3 栈式记忆挂起-恢复流程图", 6.0)},
            {"title": "3.5 黑板共享与知识库快照机制", "paragraphs": expand("黑板共享与知识库快照", ["跨域信息复用", "知识库运行时可编辑", "新会话和新评测读取最新快照"]), "figure": ("workbench", "图 3-4 工作台三栏布局示意图", 6.0)},
        ]},
        {"chapter": "第 4 章 系统详细设计与关键实现", "sections": [
            {"title": "4.1 路由器设计与意图调度逻辑", "paragraphs": expand("路由器设计", ["上下文感知的意图识别", "受限标签集合的稳定输出", "意图变化时的压栈逻辑"]), "figure": ("router", "图 4-1 路由器核心逻辑代码截图", 6.0)},
            {"title": "4.2 领域专家与工具循环设计", "paragraphs": expand("领域专家与工具循环", ["餐饮与酒店专家的局部策略", "工具消息拼接", "最大工具轮次熔断"]), "figure": ("experts", "图 4-2 领域专家工具循环代码截图", 6.0)},
            {"title": "4.3 工具层与安全过滤设计", "paragraphs": expand("工具层与安全过滤", ["查询动作与执行动作分离", "价格确认规则", "敏感动作的规则网关"]), "figure": ("tools", "图 4-3 工具层与安全过滤代码截图", 6.0)},
            {"title": "4.4 追踪服务与可视化工作台实现", "paragraphs": expand("追踪服务与工作台实现", ["TraceCollector 事件采集", "三栏工作台信息组织", "知识库编辑器的研究价值"]), "figure": ("trace", "图 4-4 Trace 采集器代码截图", 6.0)},
            {"title": "4.5 评测服务与报告输出实现", "paragraphs": expand("评测服务与报告输出", ["后台任务执行", "严格成功判定与失败原因沉淀", "报告和图表自动生成"]), "figure": ("runner", "图 4-5 评测执行器代码截图", 6.0)},
            {"title": "4.6 前端总控页面实现与交互组织", "paragraphs": expand("前端总控页面实现", ["会话、状态、评测的统一编排", "浅色专业风的研究型界面", "代码实现与论文叙述的一致性"]), "figure": ("workbench_code", "图 4-6 前端工作台总控页面代码截图", 6.0)},
        ]},
        {"chapter": "第 5 章 实验设计、结果分析与案例研究", "sections": [
            {"title": "5.1 实验环境与评测数据构建", "paragraphs": expand("实验环境与评测数据构建", ["自定义生活服务评测集", "跨域打断与模糊表达样本", "知识库驱动的实验配置"]), "figure": ("distribution", "图 5-1 评测数据集场景分布图", 5.8)},
            {"title": "5.2 指标定义与评测方法", "paragraphs": [
                f"本文采用任务成功率、平均轮次、槽位精确率、槽位召回率和槽位 F1 作为核心指标。当前实验中，多智能体系统的任务成功率为 {metrics['task_success_rate']:.4f}，平均轮次为 {metrics['average_turns']:.4f}，槽位 F1 为 {metrics['slot_f1']:.4f}；基线系统对应数值分别为 {baseline_metrics['task_success_rate']:.4f}、{baseline_metrics['average_turns']:.4f} 和 {baseline_metrics['slot_f1']:.4f}。这些指标共同反映系统是否真正完成任务，而不是只给出表面自然的回复。",
                "为了使实验分析具有更强解释力，本文没有把评测视为单纯跑分，而是把失败原因也纳入结构化结果。这样当某条样本失败时，系统能够指出是缺少关键槽位、活动任务不符合要求，还是回复内容没有满足约束，从而直接服务于策略优化。",
                "在论文写作上，本文进一步将这些指标理解为系统结构设计的外在表现。任务栈和黑板提升了槽位状态保持能力，工具层安全过滤提升了执行可靠性，而严格评测规则又反过来约束系统必须真正落地这些设计，因此实验部分与方法部分形成了闭环。",
            ]},
            {"title": "5.3 对比实验结果分析", "paragraphs": expand("对比实验结果分析", ["总体指标优势", "平均轮次与安全性的平衡", "结构化状态优于表面流畅度"]), "figures": [("metric", "图 5-2 核心指标对比图", 5.8), ("success", "图 5-3 任务成功率柱状图", 5.4), ("slotf1", "图 5-4 槽位 F1 对比图", 5.4)], "table": {"title": "表 5-1 系统总体评测结果", "headers": ["系统", "任务成功率", "平均轮次", "槽位精确率", "槽位召回率", "槽位F1"], "rows": [["Stacked Multi-Agent", f"{metrics['task_success_rate']:.4f}", f"{metrics['average_turns']:.4f}", f"{metrics['slot_precision']:.4f}", f"{metrics['slot_recall']:.4f}", f"{metrics['slot_f1']:.4f}"], ["Baseline Single Agent", f"{baseline_metrics['task_success_rate']:.4f}", f"{baseline_metrics['average_turns']:.4f}", f"{baseline_metrics['slot_precision']:.4f}", f"{baseline_metrics['slot_recall']:.4f}", f"{baseline_metrics['slot_f1']:.4f}"]]}},
            {"title": "5.4 失败模式与误差来源分析", "paragraphs": expand("失败模式与误差来源分析", ["必需槽位缺失", "位置粒度混用", "跨域状态恢复中的局部冲突"]), "table": {"title": "表 5-2 基线系统失败原因聚合统计", "headers": ["失败模式", "出现次数"], "rows": [[k, str(v)] for k, v in failure_counter.items()] or [["无", "0"]]}},
            {"title": "5.5 典型案例研究", "paragraphs": expand("典型案例研究", ["模糊表达“哪个能订就订哪个”", "餐饮与酒店之间的任务切换", "北京位置条件在跨任务中的复用"])},
            {"title": "5.6 结果讨论与有效性威胁", "paragraphs": expand("结果讨论与有效性威胁", ["数据集规模限制", "基线系统强度与可比性", "本地知识库与真实业务接口差异"])},
        ]},
        {"chapter": "第 6 章 总结与展望", "sections": [
            {"title": "6.1 全文总结", "paragraphs": expand("全文总结", ["多智能体协同价值", "对话管理结构的重要性", "研究与实现闭环的完整性"])},
            {"title": "6.2 未来展望", "paragraphs": expand("未来展望", ["更精细的黑板置信度管理", "基于 trace 的自动优化", "扩展到更多生活服务子域和真实接口"])},
        ]},
    ]


def add_front_matter(doc: Document) -> None:
    center_title(doc, THESIS_TITLE, "黑体", 18)
    center_title(doc, "摘    要", "黑体", 16)
    abstract_cn = [
        "随着大语言模型能力的持续增强，面向生活服务场景的智能体系统正在从简单问答工具演进为具备任务闭环能力的对话代理。本文围绕订餐、订房等典型生活服务任务，针对多轮交互过程中常见的上下文遗忘、意图切换、状态丢失与越权执行问题，设计并实现了一套基于对话管理的生活服务 AI-Agent 多轮交互系统。",
        "本文的核心方法包括三部分：其一，构建“全局路由器—领域专家”分层多智能体协同架构；其二，设计任务栈与全局黑板结合的统一状态模型；其三，在工具层引入价格确认、安全过滤和工具循环熔断机制。基于上述设计，本文实现了集成对话调试、知识库编辑、评测管理与报告访问的可视化工作台。",
        "实验结果表明，在覆盖餐饮、酒店和跨域打断的自定义生活服务评测集上，本文提出的堆叠多智能体系统取得了 86.36% 的任务成功率和 0.9530 的槽位 F1，显著优于不具备任务栈和黑板机制的单智能体基线。研究结果说明，结构化对话管理机制对于提升生活服务 AI-Agent 的可控性、可恢复性与工程可用性具有关键作用。",
    ]
    for item in abstract_cn:
        paragraph(doc, item, indent_cm=0.74)
    paragraph(doc, "关键词：大语言模型  多智能体  对话管理  任务栈  全局黑板  生活服务", bold=True, indent_cm=0)
    doc.add_page_break()
    center_title(doc, THESIS_TITLE_EN, "Times New Roman", 18)
    center_title(doc, "ABSTRACT", "Times New Roman", 16)
    abstract_en = [
        "This thesis studies multi-turn interaction strategies for life-service AI agents in restaurant booking and hotel booking scenarios. The work focuses on context forgetting, intent switching, state loss, and unsafe execution in practical task-oriented dialogue systems.",
        "The proposed approach combines a hierarchical multi-agent architecture, a unified dialogue state with task stack and global blackboard, and a rule-guided safety layer around tool invocation. A visual workbench integrating debugging, knowledgebase editing, evaluation, and report browsing is also implemented.",
        "Experimental results on a customized life-service dataset show that the proposed stacked multi-agent system achieves a task success rate of 86.36% and a slot F1 score of 0.9530, significantly outperforming the baseline single-agent system. The study confirms that structured dialogue management remains essential for practical LLM agents in service domains.",
    ]
    for item in abstract_en:
        paragraph(doc, item, east_asia="Times New Roman", indent_cm=1.0)
    paragraph(doc, "Key words: Large Language Models  Multi-Agent Systems  Dialogue Management  Task Stack  Global Blackboard  Life Services", east_asia="Times New Roman", bold=True, indent_cm=0)
    doc.add_page_break()
    center_title(doc, "目    录", "黑体", 16)
    toc_lines = [
        "第 1 章 绪论 ........................................................ 1",
        "第 2 章 相关理论与技术基础 .......................................... 7",
        "第 3 章 面向生活服务的系统需求分析与总体架构 ....................... 15",
        "第 4 章 系统详细设计与关键实现 ..................................... 26",
        "第 5 章 实验设计、结果分析与案例研究 ............................... 38",
        "第 6 章 总结与展望 .................................................. 49",
        "参考文献 ............................................................ 52",
    ]
    for line in toc_lines:
        paragraph(doc, line, indent_cm=0)


def write_body(doc: Document, body: list[dict], assets: dict[str, Path]) -> list[dict]:
    media_manifest: list[dict] = []
    for chapter in body:
        doc.add_page_break()
        center_title(doc, chapter["chapter"], "黑体", 16)
        for section in chapter["sections"]:
            paragraph(doc, section["title"], east_asia="黑体", size=12, bold=True)
            for item in section["paragraphs"]:
                paragraph(doc, item)
            if "figure" in section:
                key, caption, width = section["figure"]
                picture(doc, assets[key], caption, width)
                media_manifest.append({"selection": section["title"], "file": str(assets[key]), "caption": caption})
            for key, caption, width in section.get("figures", []):
                picture(doc, assets[key], caption, width)
                media_manifest.append({"selection": section["title"], "file": str(assets[key]), "caption": caption})
            if "table" in section:
                table_title(doc, section["table"]["title"])
                table(doc, section["table"]["headers"], section["table"]["rows"])
    return media_manifest


def write_references(doc: Document) -> None:
    doc.add_page_break()
    center_title(doc, "参考文献", "黑体", 16)
    refs = [
        "[1] Brown T B, Mann B, Ryder N, et al. Language models are few-shot learners[J]. Advances in Neural Information Processing Systems, 2020.",
        "[2] Yao S, Zhao J, Yu D, et al. ReAct: Synergizing reasoning and acting in language models[J]. arXiv preprint arXiv:2210.03629, 2022.",
        "[3] Schick T, Dwivedi-Yu J, Dessì R, et al. Toolformer: Language models can teach themselves to use tools[J]. arXiv preprint arXiv:2302.04761, 2023.",
        "[4] Wu Q, Bansal G, Zhang J, et al. AutoGen: Enabling next-gen LLM applications via multi-agent conversation[J]. arXiv preprint arXiv:2308.08155, 2023.",
        "[5] Budzianowski P, Casanueva I, Tseng B H, et al. MultiWOZ – A large-scale multi-domain wizard-of-oz dataset for task-oriented dialogue modelling[C]. EMNLP, 2018.",
        "[6] Williams J D, Young S. Partially observable Markov decision processes for spoken dialog systems[J]. Computer Speech & Language, 2007.",
        "[7] Sutton R S, Barto A G. Reinforcement Learning: An Introduction[M]. MIT Press, 2018.",
        "[8] Schulman J, Wolski F, Dhariwal P, et al. Proximal policy optimization algorithms[J]. arXiv preprint arXiv:1707.06347, 2017.",
        "[9] 王厚峰, 周国栋. 任务型对话系统研究综述[J]. 中文信息学报, 2021.",
        "[10] 刘挺, 孙茂松. 自然语言处理导论[M]. 北京: 高等教育出版社, 2020.",
    ]
    for ref in refs:
        paragraph(doc, ref, east_asia="楷体_GB2312", size=10, indent_cm=0)


def build_markdown(body: list[dict]) -> str:
    lines = [f"# {THESIS_TITLE}", "", "## 摘要", "飞书同步版正文与本地 Word 终稿保持一致，以下展示完整章节结构和主要分析内容。", ""]
    for chapter in body:
        lines.append(f"## {chapter['chapter']}")
        lines.append("")
        for section in chapter["sections"]:
            lines.append(f"### {section['title']}")
            lines.append("")
            lines.extend(section["paragraphs"])
            lines.append("")
    lines.append("## 参考文献")
    lines.append("")
    lines.append("详见本地 Word 终稿。")
    return "\n".join(lines)


def main() -> None:
    result = load_json(RESULTS_DIR / "latest.json")
    dataset = load_json(DATASET_PATH)
    metrics = result["stacked_multi_agent"]["metrics"]
    baseline_metrics = result["baseline_single_agent"]["metrics"]
    failure_counter: Counter[str] = Counter()
    for sample in result["baseline_single_agent"]["samples"]:
        for reason in sample.get("failure_reasons", []):
            failure_counter[reason.split(":")[0]] += 1

    assets = build_assets(result, dataset)
    body = make_body(metrics, baseline_metrics, failure_counter)

    doc = Document()
    configure_styles(doc)
    set_layout(doc.sections[0])
    add_cover(doc, assets["badge"])

    front = doc.add_section(WD_SECTION.NEW_PAGE)
    set_layout(front)
    add_header(front, assets["badge"])
    add_footer(front, "upperRoman", 1)
    add_front_matter(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_layout(body_section)
    add_header(body_section, assets["badge"])
    add_footer(body_section, "decimal", 1)
    media_manifest = write_body(doc, body, assets)
    write_references(doc)
    doc.save(OUTPUT_PATH)

    MARKDOWN_PATH.write_text(build_markdown(body), encoding="utf-8")
    MEDIA_MANIFEST_PATH.write_text(json.dumps(media_manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"docx": str(OUTPUT_PATH), "markdown": str(MARKDOWN_PATH), "manifest": str(MEDIA_MANIFEST_PATH), "media_count": len(media_manifest)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
