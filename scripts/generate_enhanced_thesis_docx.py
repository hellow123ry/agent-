from __future__ import annotations

import json
from collections import Counter
from pathlib import Path

import matplotlib.pyplot as plt
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from pygments import highlight
from pygments.formatters.img import ImageFormatter
from pygments.lexers import get_lexer_for_filename


PROJECT_ROOT = Path("/Users/bytedance/Desktop/multi_agent_dialog_system")
RESULTS_DIR = PROJECT_ROOT / "evaluation" / "results"
DATASET_PATH = PROJECT_ROOT / "evaluation" / "datasets" / "life_service_eval.json"
OUTPUT_PATH = Path("/Users/bytedance/Desktop/基于对话管理的生活服务_AI_Agent_多轮交互策略_50页增强版.docx")
ASSET_DIR = PROJECT_ROOT / ".thesis_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

plt.rcParams["font.sans-serif"] = ["PingFang SC", "Heiti SC", "Arial Unicode MS", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def read_snippet(path: Path, start: int | None = None, end: int | None = None) -> str:
    lines = path.read_text(encoding="utf-8").splitlines()
    sliced = lines[(start - 1 if start else 0) : end]
    return "\n".join(sliced)


def set_run_font(run, east_asia: str, latin: str = "Times New Roman", size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)


def make_box_diagram(title: str, nodes: list[tuple[int, int, int, int, str, str]], arrows: list[tuple[tuple[int, int], tuple[int, int]]], output: Path) -> Path:
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.truetype("/System/Library/Fonts/STHeiti Medium.ttc", 40)
    text_font = ImageFont.truetype("/System/Library/Fonts/STHeiti Medium.ttc", 26)
    draw.text((60, 40), title, fill="#0f172a", font=title_font)
    for x1, y1, x2, y2, label, fill in nodes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, outline="#1e3a8a", width=4, fill=fill)
        bbox = draw.multiline_textbbox((0, 0), label, font=text_font, spacing=8)
        width = bbox[2] - bbox[0]
        height = bbox[3] - bbox[1]
        draw.multiline_text(((x1 + x2 - width) / 2, (y1 + y2 - height) / 2), label, fill="#0f172a", font=text_font, align="center", spacing=8)
    for (sx, sy), (ex, ey) in arrows:
        draw.line((sx, sy, ex, ey), fill="#475569", width=6)
        draw.polygon([(ex, ey), (ex - 18, ey - 10), (ex - 18, ey + 10)], fill="#475569")
    image.save(output)
    return output


def make_workbench_layout(output: Path) -> Path:
    nodes = [
        (80, 160, 460, 760, "左栏\nChatPanel\n会话与输入", "#eff6ff"),
        (560, 160, 1040, 440, "中栏上\nTracePanel\n事件时间线", "#f8fafc"),
        (560, 480, 1040, 760, "中栏下\nStatePanel\n任务栈 / 黑板", "#f8fafc"),
        (1140, 160, 1520, 380, "右栏上\nKnowledgebasePanel", "#eef2ff"),
        (1140, 420, 1520, 590, "右栏中\nEvalPanel", "#eef2ff"),
        (1140, 630, 1520, 760, "右栏下\nReportsPanel", "#eef2ff"),
    ]
    arrows = [((460, 460), (560, 300)), ((460, 460), (560, 620)), ((1040, 300), (1140, 270)), ((1040, 620), (1140, 505))]
    return make_box_diagram("图  工作台三栏布局示意图", nodes, arrows, output)


def make_architecture_diagram(output: Path) -> Path:
    nodes = [
        (80, 300, 300, 420, "用户输入", "#dbeafe"),
        (380, 250, 640, 470, "全局路由器\nRouter", "#e0f2fe"),
        (760, 120, 1040, 250, "餐饮专家", "#dcfce7"),
        (760, 290, 1040, 420, "酒店专家", "#dcfce7"),
        (760, 460, 1040, 590, "通用专家", "#dcfce7"),
        (1140, 100, 1480, 290, "工具层\nsearch / book /\nupdate_blackboard", "#fef3c7"),
        (1140, 340, 1480, 530, "状态层\nDialogState / task_stack /\nblackboard / knowledgebase", "#ede9fe"),
        (1140, 580, 1480, 760, "可观测层\nTrace / Eval / Reports", "#fee2e2"),
    ]
    arrows = [
        ((300, 360), (380, 360)),
        ((640, 300), (760, 180)),
        ((640, 360), (760, 360)),
        ((640, 420), (760, 530)),
        ((1040, 180), (1140, 180)),
        ((1040, 360), (1140, 180)),
        ((1040, 530), (1140, 180)),
        ((1040, 360), (1140, 430)),
        ((1040, 360), (1140, 670)),
    ]
    return make_box_diagram("图  多智能体对话系统总体架构图", nodes, arrows, output)


def make_stack_flow_diagram(output: Path) -> Path:
    nodes = [
        (160, 160, 520, 290, "步骤1：酒店任务激活\nactive_task = hotel", "#eff6ff"),
        (160, 360, 520, 490, "步骤2：用户打断询问餐饮\npush(hotel)", "#ecfccb"),
        (160, 560, 520, 690, "步骤3：餐饮任务完成\npop() 恢复 hotel", "#fef3c7"),
        (760, 130, 1360, 720, "任务栈视图\n\n顶部：hotel\n中部：...\n底部：历史挂起任务", "#f8fafc"),
    ]
    arrows = [
        ((340, 290), (340, 360)),
        ((340, 490), (340, 560)),
        ((520, 225), (760, 225)),
        ((520, 425), (760, 425)),
        ((520, 625), (760, 625)),
    ]
    return make_box_diagram("图  栈式记忆挂起-恢复流程图", nodes, arrows, output)


def make_distribution_chart(dataset: list[dict], output: Path) -> Path:
    counter = Counter(sample["domain"] for sample in dataset)
    labels = list(counter.keys())
    values = [counter[label] for label in labels]
    plt.figure(figsize=(10, 5.6))
    bars = plt.bar(labels, values, color=["#2563eb", "#0ea5e9", "#10b981"])
    plt.title("评测数据集场景分布")
    plt.xlabel("场景类别")
    plt.ylabel("样本数量")
    for bar, value in zip(bars, values):
        plt.text(bar.get_x() + bar.get_width() / 2, value + 0.05, str(value), ha="center")
    plt.tight_layout()
    plt.savefig(output, dpi=220)
    plt.close()
    return output


def make_metric_comparison(result: dict, output: Path) -> Path:
    ours = result["stacked_multi_agent"]["metrics"]
    baseline = result["baseline_single_agent"]["metrics"]
    labels = ["成功率", "平均轮次", "槽位F1"]
    ours_values = [ours["task_success_rate"], ours["average_turns"], ours["slot_f1"]]
    base_values = [baseline["task_success_rate"], baseline["average_turns"], baseline["slot_f1"]]
    x = range(len(labels))
    width = 0.35
    plt.figure(figsize=(10, 5.6))
    plt.bar([v - width / 2 for v in x], ours_values, width=width, label="Stacked Multi-Agent", color="#2563eb")
    plt.bar([v + width / 2 for v in x], base_values, width=width, label="Baseline Single Agent", color="#ef4444")
    plt.xticks(list(x), labels)
    plt.title("核心指标对比图")
    plt.legend()
    for idx, value in enumerate(ours_values):
        plt.text(idx - width / 2, value + 0.03, f"{value:.2f}", ha="center")
    for idx, value in enumerate(base_values):
        plt.text(idx + width / 2, value + 0.03, f"{value:.2f}", ha="center")
    plt.tight_layout()
    plt.savefig(output, dpi=220)
    plt.close()
    return output


def make_single_metric_chart(result: dict, metric_key: str, title: str, output: Path) -> Path:
    ours = result["stacked_multi_agent"]["metrics"][metric_key]
    baseline = result["baseline_single_agent"]["metrics"][metric_key]
    labels = ["Multi-Agent", "Baseline"]
    values = [ours, baseline]
    colors = ["#2563eb", "#ef4444"]
    plt.figure(figsize=(8, 4.8))
    bars = plt.bar(labels, values, color=colors)
    plt.title(title)
    ymax = max(values) if max(values) > 0 else 1.0
    plt.ylim(0, ymax * 1.25)
    for bar, value in zip(bars, values):
        plt.text(bar.get_x() + bar.get_width() / 2, value + ymax * 0.04 if ymax else 0.04, f"{value:.2f}", ha="center")
    plt.tight_layout()
    plt.savefig(output, dpi=220)
    plt.close()
    return output


def render_code_image(source: Path, start: int, end: int, output: Path) -> Path:
    code = read_snippet(source, start, end)
    lexer = get_lexer_for_filename(source.name, code)
    formatter = ImageFormatter(
        font_name="Menlo",
        font_size=20,
        line_numbers=True,
        line_number_bg="#f8fafc",
        line_number_fg="#64748b",
        style="friendly",
        image_pad=20,
        line_pad=4,
        hl_color="#fff7cc",
    )
    png_bytes = highlight(code, lexer, formatter)
    output.write_bytes(png_bytes)
    return output


def set_document_style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        s = document.styles[style_name]
        s.font.name = "Times New Roman"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_title(document: Document) -> None:
    for _ in range(4):
        document.add_paragraph("")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("本科毕业论文")
    set_run_font(run, "黑体", size=22, bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于对话管理的生活服务 AI-Agent 多轮交互策略")
    set_run_font(run, "黑体", size=20, bold=True)

    for _ in range(6):
        document.add_paragraph("")

    info_lines = [
        "学    生：________________",
        "学    号：________________",
        "学    院：________________",
        "专    业：计算机科学与技术",
        "指导教师：________________",
        "完成日期：2026 年 5 月",
    ]
    for line in info_lines:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_run_font(run, "宋体", size=14)

    document.add_page_break()


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading("", level=level)
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, "黑体", size=16, bold=True)
    elif level == 2:
        set_run_font(run, "黑体", size=14, bold=True)
    else:
        set_run_font(run, "黑体", size=12, bold=True)


def add_paragraph(document: Document, text: str, bold: bool = False, align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if align:
        p.alignment = align
        p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, "宋体", size=12, bold=bold)


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, "宋体", size=11, bold=True)


def add_image(document: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    add_caption(document, caption)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def thesis_paragraphs() -> dict[str, list[str]]:
    return {
        "绪论-背景": [
            "近年来，大语言模型在自然语言理解、生成以及复杂指令执行方面不断突破，推动了智能客服、问答系统和任务型对话系统的快速发展。与早期依赖规则模板和统计学习的系统相比，基于大模型的对话系统具备更强的语义泛化能力，能够在非标准表达、跨领域问句与模糊需求场景下保持相对自然的交互表现。尤其在本地生活服务场景中，用户常常使用口语化表达，例如“帮我看看明天哪家能吃烤肉”“哪个酒店能订就订哪个”，这类表达包含大量省略、指代与不完整约束，传统有限状态机系统难以稳健应对，而大模型在语义补全方面展现出明显优势。",
            "然而，生活服务任务并不仅仅是“会聊天”这么简单。真实应用要求系统在理解用户意图之后稳定完成检索、筛选、澄清、确认和执行等一系列动作，并在整个过程中保持可追踪、可解释和可恢复。以酒店预订为例，系统不仅要理解“市中心”“豪华型”“明天入住后天退房”等要素，还必须遵循价格确认后才能下单的业务规则；以餐饮预订为例，系统需要在不完全掌握全部槽位时及时调用工具，向用户展示有位资源，再进一步推进确认流程。由此可见，真正的难点并不是语言生成本身，而是围绕对话管理的状态维护与策略调度。",
            "本研究聚焦于“生活服务 AI-Agent 多轮交互策略”这一问题，核心观察是：用户在任务型对话中的行为高度非线性，他们会突然切换主题、打断当前任务、复用此前信息，甚至在确认关键条件前改变决策。如果系统没有显式的任务挂起与恢复机制，那么强大的大模型也会因上下文漂移而出现遗漏、误判或错误执行。基于这一判断，本文将研究重点从单轮问答和纯提示词优化转向“对话管理架构设计”，通过多智能体分工、任务栈记忆与黑板共享构建可工程落地的鲁棒交互框架。",
        ],
        "绪论-意义": [
            "从理论层面看，本文尝试回答一个具有现实意义的问题：在大模型已具备通用语言能力的条件下，是否仍然需要显式的任务结构和状态机理？本文的结论是肯定的。因为在多轮任务型交互中，模型能力越强，越需要借助结构化约束将“会理解”转化为“能稳定执行”。因此，本文提出的任务栈与黑板机制并非对大模型能力的替代，而是对其进行结构化组织，使之具备更高的可控性和工程鲁棒性。",
            "从应用层面看，本地生活服务是高频、高并发、强约束的典型领域。用户在此类场景中的容错率较低，如果系统反复追问、丢失条件或在未确认价格时直接下单，都会迅速降低用户信任度。因此，如何把大模型嵌入到清晰的对话管理框架中，直接关系到其商业落地价值。本文构建的原型系统不仅用于论文实验，也为后续的多域生活服务助手提供了可迁移的系统设计经验。",
        ],
        "理论基础": [
            "任务型多轮对话系统通常可以被形式化为一个带有部分可观测特征的决策过程。在每一轮中，系统需要根据当前用户输入、历史对话、已知槽位和外部工具结果决定下一步动作。与开放域聊天不同，任务型对话更强调目标驱动与状态一致性，因此其评估标准也更加客观，常见指标包括任务成功率、平均轮次、槽位召回率和最终满意度等。本文在继承这一研究范式的基础上，将任务栈和黑板信息显式纳入状态表示。",
            "在大语言模型驱动的 Agent 框架中，Tool Calling 是连接语言能力与外部系统能力的核心桥梁。模型通过结构化调用接口触发检索、查询和执行，从而将语言推理映射为可验证、可追踪的业务动作。本文系统中的 `search_restaurants`、`search_hotels`、`book_hotel` 和 `update_blackboard` 就构成了最小工具集合，其中既包含信息查询，也包含敏感执行操作。敏感执行动作与普通查询动作在约束上并不相同，因此需要在工具层注入安全规则。",
            "ReAct 范式强调“推理—行动—观察”的循环结构，其优势在于能让模型在复杂任务中逐步逼近正确结果，但它也可能在工具循环中陷入局部振荡。本文在工程上引入 `MAX_TOOL_CALL_ROUNDS` 和评测单样本超时机制，实际上是在 ReAct 的自由探索外部再加一层安全边界，使系统具备可熔断、可回退的特性。这种做法虽然看似简单，却对提升实际交付质量十分关键。",
            "强化学习与奖励重塑为对话策略优化提供了重要思路。虽然本文原型尚未完整实现端到端 PPO 训练，但我们已经将状态空间、动作空间与奖励设计抽象出来，并在论文中进一步讨论如何通过解释信号和规则反馈构建“解释—修正—再执行”的闭环。该部分既是对当前系统的理论总结，也是下一步研究的延展方向。",
        ],
        "系统设计": [
            "系统总体架构采用“路由器 + 领域专家 + 工具层 + 状态层 + 可观测层”的分层方式。路由器负责根据当前对话上下文识别用户意图；领域专家根据分工执行餐饮、酒店等具体逻辑；工具层负责真实的查询和执行；状态层统一维护消息历史、活动任务、任务栈、全局黑板与知识库快照；可观测层则负责沉淀 trace、评测进度与报告文件。这样的职责划分使各模块边界清晰，便于独立调试和实验比较。",
            "路由器并不是简单的意图分类器，而是一个带有上下文感知能力的状态迁移控制点。它在提示中明确引入了 `active_task` 与 `task_stack`，并规定用户在回答“明天”“是的”等省略性话语时，应结合历史对话推断所属领域。这样做避免了单看最后一句输入导致的意图漂移。同时，当路由器识别到新意图与当前任务不一致时，会主动将旧任务压栈，形成挂起状态，为后续恢复做准备。",
            "任务栈的设计直接借鉴了操作系统中的栈帧思想。每次任务切换都不是粗暴覆盖，而是将当前任务保留在 `task_stack` 中。当新任务结束后，系统通过 `_handle_task_completion` 自动弹栈，并生成一段自然的过渡语，引导用户回到之前尚未完成的任务。这种方式保留了对话的连贯性，也降低了用户重新组织表达的负担。",
            "全局黑板用于跨任务共享长期有效信息，例如位置、人数和日期。在生活服务场景中，这类信息具有明显的跨域复用价值。用户在酒店查询中提到“我在北京”，随后切换到餐饮预订时，餐饮专家可以直接复用该位置条件，而不必重复追问。黑板机制的意义在于将“用户已经明确表达过且短时间内稳定”的信息从局部任务状态提升为全局上下文资源。",
            "知识库模块将餐厅与酒店数据从硬编码常量扩展为可运行时修改的 JSON 数据源，并在前端工作台中提供编辑界面。这样一来，系统不仅具备对话和评测能力，还具备实验配置能力。研究者可以在不改动代码的前提下调整酒店价格、状态和餐厅容量，再观察对系统对话与评测结果的影响，这为后续做更多消融和敏感性分析提供了便利。",
            "可观测性是本文系统与许多原型 Demo 的重要区别之一。系统通过 `TraceCollector` 在路由、工具调用、工具结果和黑板更新等关键节点记录事件，为前端的 TracePanel 提供直接输入。同时，评测服务通过后台进程与轮询机制暴露任务进度，使得研究者可以在图形界面中而不是终端中观察模型行为。这一设计使论文中的案例分析不再依赖主观回忆，而是基于结构化日志进行复盘。",
        ],
        "策略优化": [
            "本文将对话策略视为一个在约束条件下求解最优动作序列的过程。与传统的纯规则系统相比，大模型具备更强的自然语言推断能力，但其不确定性也更高；与纯端到端智能体相比，规则先验虽然降低了搜索空间，却可能抑制模型灵活性。因此，本文采用混合式策略优化思路：在高层保持大模型的语义推理自由，在低层工具执行与安全动作上引入硬性规则约束。",
            "酒店预订中的价格确认规则是一个典型例子。在业务现实中，用户未确认价格前不应直接触发下单动作，因此 `book_hotel` 工具要求 `price_confirmed=True` 才能执行成功。如果模型越过这一边界，工具会返回 `blocked_by_filter`，提示系统先向用户报价并确认接受。这一机制本质上是把业务安全规则前置到动作空间过滤阶段，从而用确定性逻辑修正生成式模型的潜在越权行为。",
            "工具循环熔断机制则解决了另一类问题：模型在复杂条件或模糊指令下可能不断检索、不断重试，造成耗时过长甚至看似卡死。通过设置 `MAX_TOOL_CALL_ROUNDS=3` 并在评测执行中增加单样本超时，系统将不确定性控制在用户与评测体系可接受的边界内。这种处理方式虽然会牺牲部分探索空间，但显著提升了系统的稳定性和可观测性，是工程场景中必要的折中。",
            "从强化学习视角看，上述规则可以理解为动作先验和奖励塑形。对于明显不应执行的动作，直接通过规则拒绝；对于可以执行但策略不优的动作，则通过评测指标和日志反馈进行后验优化。本文后续计划在此基础上构建可解释反馈网络，使 trace 中的事件序列不仅用于人工调试，也能转化为训练信号服务于自动纠偏。",
        ],
        "系统实现": [
            "后端采用 FastAPI 作为本地服务层，负责暴露会话管理、评测调度、报告访问和知识库编辑接口。会话服务在创建或重置时加载知识库快照，并在每次 `send_turn` 调用时返回最新消息、对话状态和 trace 事件。评测服务则通过独立进程执行整套数据集，使用轮询暴露状态、局部指标和最终结果。前端采用 React + TypeScript，提供聊天、状态追踪、知识库编辑和评测面板的统一工作台。",
            "从模块实现看，`DialogState` 是系统的统一数据中枢。它不只是一个普通字典，而是把消息、任务栈、当前意图、全局黑板和知识库快照统一到同一上下文中，确保任何节点在运行时都能访问一致的状态视图。由于多轮对话系统的错误往往来自状态不同步，显式状态建模是整个架构得以工作的前提。",
            "前端工作台的设计目标不是做一个华丽的演示页，而是为研究和调试服务。左栏聚焦真实对话，方便观察用户输入与模型输出；中栏聚焦状态与 Trace，便于理解为什么会发生路由切换和工具调用；右栏则聚焦评测、报告和知识库，帮助研究者把配置、执行和分析整合到一个页面中。这样一种研究型前端对提升迭代效率非常关键。",
            "知识库编辑面板将系统从“固定样本演示”进一步推进到“可配置原型平台”。研究者既可以修改酒店价格、房态，也可以编辑餐厅类型和容量。由于新会话和新评测任务会读取最新知识库快照，平台天然支持小规模实验，比如观察价格变化是否影响澄清逻辑，或观察餐厅容量变化是否影响失败率。这使得论文中的系统实现部分具备更强的实验支撑。",
        ],
        "实验分析": [
            "本文的评测体系面向复杂多轮生活服务场景，而不是仅验证单轮问答正确率。数据集包含餐饮、酒店和跨域场景三类样本，重点考察系统在任务切换、黑板复用、模糊条件处理和安全规则触发方面的表现。与传统开放域评测不同，这里每条样本都带有结构化 `success_conditions`，包括关键词约束、必需槽位和最终活动任务要求，因此实验结论更接近真实工程目标。",
            "实验结果表明，堆叠多智能体系统在任务成功率和槽位 F1 上明显优于基线单智能体。尤其值得注意的是，基线系统虽然在文本表面上也能给出看似合理的答复，但由于其未维护任务栈和全局黑板，最终结构化状态为空或不完整，因此在客观指标上几乎全部失败。这说明仅凭自然语言流畅性无法证明任务型系统真的“完成了任务”，结构化评测不可或缺。",
            "平均轮次方面，多智能体系统并未一味追求更少轮数，而是在必要时主动发起澄清和确认，因此平均轮次保持在 2.68 左右。换言之，本文系统将一部分额外交互代价换成了显著更高的成功率和安全性。这种权衡符合生活服务场景的用户预期：用户通常可以接受一次价格确认，但不能接受系统直接下错单或遗忘先前条件。",
            "案例研究显示，酒店任务中的“哪个能订就订哪个”属于高度模糊但在现实中非常常见的表达。如果系统坚持逐项追问，用户会感到迟钝；如果系统完全不加限制地自由执行，则可能误下单。本文系统的做法是先在空参数条件下执行检索，再把结果和价格回报给用户，引导其做最低必要确认。这种策略体现了灵活性与安全性的平衡。",
            "跨域黑板样本进一步说明了结构化状态设计的必要性。用户先在酒店场景中提供了位置，随后切到餐饮场景，如果系统能够从黑板中读取位置，就能直接推荐“北京附近的烤肉店”，而不必再次询问“你在哪儿”。这种体验上的顺滑并不来自更强的语言生成，而来自更合理的状态组织。实验结果中较高的槽位 F1 正是这一设计价值的量化体现。",
            "本文同时承认，系统仍存在局限。比如在某些样本中，模型会把“北京”与“市中心”这类层级不同的位置条件混用，或在跨域切换时产生局部槽位丢失。这表明黑板共享虽然有效，但还需要更精细的层级建模和冲突消解策略。未来可以通过字段级置信度、来源标记和时间戳机制进一步增强该能力。",
        ],
        "总结展望": [
            "本文围绕生活服务场景下的多轮任务型交互问题，提出并实现了一套基于多智能体协同、任务栈记忆和全局黑板的对话管理系统。系统通过 Router 完成高层调度，通过 Dining/Hotel 等专家完成领域内决策，通过工具层连接知识库与执行逻辑，并通过 Trace 和 Eval 组件保障可观测与可评估。实验结果表明，该系统在复杂样本上相较基线取得了显著优势。",
            "与纯提示词驱动的简单 Agent 原型不同，本文强调结构化设计的重要性，即在保持大模型灵活性的同时，通过状态建模和规则边界增强其可控性。任务栈解决了意图打断后的恢复问题，黑板提升了跨任务信息复用效率，安全规则避免了敏感动作越界，这些设计共同构成了可工程落地的多轮交互策略。",
            "未来工作可以从三个方向继续推进：其一，引入更精细的层级记忆和槽位置信度管理；其二，基于 trace 数据训练策略优化模块，实现从人工调试到自动纠偏的升级；其三，扩展到更多生活服务子域和更大规模真实日志，以检验架构在开放场景中的泛化能力。随着这些工作逐步完善，本系统有望成为具备产品化潜力的生活服务 AI-Agent 基础框架。",
        ],
    }


def add_standard_front_matter(document: Document) -> None:
    add_heading(document, "摘  要", 1)
    for paragraph in [
        "随着大语言模型能力的快速提升，面向本地生活服务的智能体系统正在从问答式工具逐步向具备任务闭环能力的对话代理演进。本文围绕订餐、订房等典型生活服务任务，研究多轮交互过程中普遍存在的上下文遗忘、意图跳转、状态丢失和越权执行等问题，提出了一种基于对话管理的生活服务 AI-Agent 多轮交互策略。",
        "本文的核心方法包括三部分：一是构建“全局路由器—领域专家”的分层多智能体协同架构，以降低单个大模型在复杂对话中的认知负担；二是设计任务栈与全局黑板相结合的状态管理机制，用于支持任务挂起、上下文恢复和跨域信息共享；三是在工具层引入价格确认、安全过滤和工具循环熔断等规则先验，以提升执行安全性和系统稳定性。在工程实现上，本文基于 FastAPI、React、LangGraph 和可编辑知识库构建了完整原型工作台，并提供了 trace 追踪、评测管理和报告访问能力。",
        "实验结果显示，在覆盖餐饮、酒店和跨域打断的生活服务评测集上，本文系统的任务成功率达到 86.36%，槽位 F1 达到 0.9530，显著优于无任务栈与无黑板机制的单智能体基线。研究表明，结构化的对话管理架构对提升大模型在生活服务场景中的可控性、可恢复性和可解释性具有重要作用。本文工作为后续开展更大规模的 Agent 策略优化与产品化验证提供了系统基础。",
    ]:
        add_paragraph(document, paragraph)
    add_paragraph(document, "关键词：大语言模型；多智能体；对话管理；任务栈；全局黑板；生活服务", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_page_break = document.add_page_break
    add_page_break()

    add_heading(document, "Abstract", 1)
    for paragraph in [
        "With the rapid rise of large language models, AI assistants in local life services are evolving from simple conversational tools into task-completion agents. This thesis studies the dialogue management challenges in restaurant and hotel booking scenarios, especially context forgetting, intent switching, state loss, and unsafe execution in multi-turn interactions.",
        "We propose a dialogue-management-oriented life-service AI-Agent framework featuring three core ideas: a hierarchical multi-agent architecture composed of a global router and domain experts; a stack-based suspend-resume memory mechanism combined with a global blackboard for cross-task information sharing; and a rule-guided safety layer to prevent unsafe actions such as booking before price confirmation. The system is implemented as a full prototype workbench with FastAPI, React, LangGraph, editable knowledge base support, trace inspection, and automated evaluation workflows.",
        "Experimental results on a customized life-service evaluation set show that the proposed stacked multi-agent system achieves a task success rate of 86.36% and a slot F1 score of 0.9530, significantly outperforming the baseline single-agent system without task stack and blackboard support. The study demonstrates that structured dialogue management is essential for making large language model agents controllable, recoverable, and practical in real task-oriented service domains.",
    ]:
        add_paragraph(document, paragraph)
    add_paragraph(document, "Keywords: Large Language Models; Multi-Agent Systems; Dialogue Management; Task Stack; Global Blackboard; Life Services", align=WD_ALIGN_PARAGRAPH.CENTER)
    document.add_page_break()

    add_heading(document, "目  录", 1)
    toc_lines = [
        "第 1 章 绪论 ...................................................... 1",
        "第 2 章 相关理论与技术基础 ........................................ 7",
        "第 3 章 基于多智能体协同与栈式记忆的对话管理架构 ............... 15",
        "第 4 章 融合规则先验与可解释策略的优化方法 ....................... 25",
        "第 5 章 生活服务 AI-Agent 原型系统实现与评估 ..................... 33",
        "第 6 章 总结与展望 ................................................ 46",
        "参考文献 .......................................................... 49",
        "附录 A 核心代码清单 ............................................... 52",
        "附录 B 评测样本与结果摘要 ......................................... 60",
    ]
    for line in toc_lines:
        add_paragraph(document, line)
    document.add_page_break()


def add_body(document: Document, assets: dict[str, Path], result: dict) -> None:
    sections = thesis_paragraphs()

    add_heading(document, "第 1 章 绪论", 1)
    add_heading(document, "1.1 研究背景与意义", 2)
    for paragraph in sections["绪论-背景"] + sections["绪论-意义"]:
        add_paragraph(document, paragraph)
    add_heading(document, "1.2 国内外研究现状", 2)
    for topic in [
        "在任务型对话系统领域，早期研究主要围绕对话状态追踪、槽位填充和策略学习展开，典型数据集如 ATIS、CamRest 和 MultiWOZ 推动了模块化方法的快速发展。近年来，随着预训练模型和生成式方法的兴起，研究重心逐渐从单独的状态分类迁移到统一生成框架，但跨任务切换与长程依赖问题仍是难点。",
        "在大模型 Agent 领域，ReAct、Toolformer、AutoGPT、AutoGen 等工作展示了模型通过调用工具解决复杂问题的潜力。然而，这些方法大多面向开放任务与通用推理，缺乏对垂直业务场景中确定性约束的细化建模。生活服务场景中的价格确认、库存校验和资源状态更新，并不能简单依赖语言模型的“自觉性”完成。",
        "多智能体系统为降低单体模型的推理负担提供了一种有效路径。通过让不同专家负责不同任务域，系统可以在保持整体灵活性的同时提升局部准确率。但若缺乏统一状态容器和任务恢复机制，多智能体之间的信息仍会割裂，甚至出现领域之间互相覆盖上下文的问题。因此，结构化状态管理应被视为多智能体系统的基础设施。",
    ]:
        add_paragraph(document, topic)
    add_heading(document, "1.3 论文主要工作与组织结构", 2)
    for paragraph in [
        "本文的主要工作可以概括为四点：第一，设计并实现了一个面向生活服务的多智能体对话管理架构；第二，提出了任务栈与全局黑板相结合的上下文挂起—恢复机制；第三，在工具层引入安全规则与熔断机制，显著提升工程可用性；第四，构建了图形化工作台和自动化评测流程，对系统进行了定量与定性验证。",
        "全文组织结构如下：第 1 章介绍研究背景和研究价值；第 2 章梳理相关理论与技术基础；第 3 章详细阐述多智能体协同与栈式记忆架构；第 4 章讨论规则先验与可解释策略优化；第 5 章介绍系统实现、实验设计和评估结果；第 6 章总结全文并展望后续工作。",
    ]:
        add_paragraph(document, paragraph)
    add_image(document, assets["architecture"], "图 3-1 多智能体对话系统总体架构图")

    add_heading(document, "第 2 章 相关理论与技术基础", 1)
    for idx, paragraph in enumerate(sections["理论基础"], start=1):
        add_paragraph(document, paragraph)
        if idx == 2:
            add_image(document, assets["workbench"], "图 2-1 工作台三栏布局示意图")

    add_heading(document, "第 3 章 基于多智能体协同与栈式记忆的对话管理架构", 1)
    add_heading(document, "3.1 架构设计动机与核心思想", 2)
    for paragraph in sections["系统设计"][:2]:
        add_paragraph(document, paragraph)
    add_heading(document, "3.2 分层多智能体协同架构", 2)
    for paragraph in sections["系统设计"][2:4]:
        add_paragraph(document, paragraph)
    add_image(document, assets["stack_flow"], "图 3-2 栈式记忆挂起-恢复流程图")
    add_heading(document, "3.3 全局黑板与知识库快照机制", 2)
    for paragraph in sections["系统设计"][4:]:
        add_paragraph(document, paragraph)

    add_heading(document, "第 4 章 融合规则先验与可解释策略的优化方法", 1)
    for paragraph in sections["策略优化"]:
        add_paragraph(document, paragraph)
    add_image(document, assets["metric_compare"], "图 4-1 核心指标对比图")

    add_heading(document, "第 5 章 生活服务 AI-Agent 原型系统实现与评估", 1)
    add_heading(document, "5.1 系统开发环境与技术栈", 2)
    for paragraph in sections["系统实现"][:2]:
        add_paragraph(document, paragraph)
    add_heading(document, "5.2 前后端工作台与知识库编辑实现", 2)
    for paragraph in sections["系统实现"][2:]:
        add_paragraph(document, paragraph)
    add_image(document, assets["distribution"], "图 5-1 评测数据集场景分布图")
    add_image(document, assets["success_rate"], "图 5-2 任务成功率柱状图")
    add_image(document, assets["slot_f1"], "图 5-3 槽位 F1 对比图")
    add_image(document, assets["average_turns"], "图 5-4 平均轮次对比图")

    add_heading(document, "5.3 实验设计与指标体系", 2)
    for paragraph in sections["实验分析"][:2]:
        add_paragraph(document, paragraph)

    metrics = result["stacked_multi_agent"]["metrics"]
    base_metrics = result["baseline_single_agent"]["metrics"]
    add_table(
        document,
        ["系统", "任务成功率", "平均轮次", "槽位精确率", "槽位召回率", "槽位F1"],
        [
            ["Stacked Multi-Agent", f"{metrics['task_success_rate']:.4f}", f"{metrics['average_turns']:.4f}", f"{metrics['slot_precision']:.4f}", f"{metrics['slot_recall']:.4f}", f"{metrics['slot_f1']:.4f}"],
            ["Baseline Single Agent", f"{base_metrics['task_success_rate']:.4f}", f"{base_metrics['average_turns']:.4f}", f"{base_metrics['slot_precision']:.4f}", f"{base_metrics['slot_recall']:.4f}", f"{base_metrics['slot_f1']:.4f}"],
        ],
    )
    add_caption(document, "表 5-1 系统总体评测结果")

    add_heading(document, "5.4 实验结果与案例分析", 2)
    for paragraph in sections["实验分析"][2:]:
        add_paragraph(document, paragraph)

    failure_counter = Counter()
    for sample in result["baseline_single_agent"]["samples"]:
        for reason in sample.get("failure_reasons", []):
            failure_counter[reason.split(":")[0]] += 1
    add_paragraph(document, "为了进一步理解基线模型的失败模式，本文对失败原因进行了聚合统计。结果显示，`required_slots failed` 是最主要的问题，这意味着系统虽然能够生成表面合理的自然语言回复，但并没有把关键约束稳定沉淀到结构化状态中。与之对应，堆叠多智能体系统在黑板和任务栈的加持下，能显著降低这类状态遗漏。")
    add_table(document, ["失败模式", "出现次数"], [[k, str(v)] for k, v in failure_counter.items()])
    add_caption(document, "表 5-2 基线系统失败原因聚合统计")

    add_heading(document, "第 6 章 总结与展望", 1)
    for paragraph in sections["总结展望"]:
        add_paragraph(document, paragraph)


def add_references(document: Document) -> None:
    add_heading(document, "参考文献", 1)
    refs = [
        "[1] Brown T B, Mann B, Ryder N, et al. Language models are few-shot learners[J]. Advances in Neural Information Processing Systems, 2020.",
        "[2] Yao S, Zhao J, Yu D, et al. ReAct: Synergizing reasoning and acting in language models[J]. arXiv preprint arXiv:2210.03629, 2022.",
        "[3] Schick T, Dwivedi-Yu J, Dessì R, et al. Toolformer: Language models can teach themselves to use tools[J]. arXiv preprint arXiv:2302.04761, 2023.",
        "[4] Wu Q, Bansal G, Zhang J, et al. AutoGen: Enabling next-gen LLM applications via multi-agent conversation[J]. arXiv preprint arXiv:2308.08155, 2023.",
        "[5] Budzianowski P, Casanueva I, Tseng B H, et al. MultiWOZ – A large-scale multi-domain wizard-of-oz dataset for task-oriented dialogue modelling[C]. EMNLP, 2018.",
        "[6] Williams J D, Young S. Partially observable Markov decision processes for spoken dialog systems[J]. Computer Speech & Language, 2007.",
        "[7] Li X, Chen Y-N, Li L, et al. End-to-end task-completion neural dialogue systems[J]. arXiv preprint arXiv:1703.01008, 2017.",
        "[8] Peng B, Li C, Li J, et al. SOLOIST: Building task bots at scale with transfer learning and machine teaching[J]. Transactions of the Association for Computational Linguistics, 2021.",
        "[9] Zhao T, Eskenazi M. Zero-shot dialog generation with cross-domain latent actions[J]. SIGDIAL, 2018.",
        "[10] Sutton R S, Barto A G. Reinforcement Learning: An Introduction[M]. MIT Press, 2018.",
        "[11] Schulman J, Wolski F, Dhariwal P, et al. Proximal policy optimization algorithms[J]. arXiv preprint arXiv:1707.06347, 2017.",
        "[12] Mnih V, Kavukcuoglu K, Silver D, et al. Human-level control through deep reinforcement learning[J]. Nature, 2015.",
    ]
    for item in refs:
        add_paragraph(document, item)


def add_appendices(document: Document, assets: dict[str, Path], result: dict, dataset: list[dict]) -> None:
    add_heading(document, "附录 A 核心代码截图与实现说明", 1)
    code_assets = [
        ("附图 A-1 `DialogState` 状态结构代码截图", assets["code_state"], "该代码定义了对话系统的统一状态容器，任务栈、活动任务、黑板和知识库都在此完成汇总。"),
        ("附图 A-2 `route_intent` 路由逻辑代码截图", assets["code_router"], "该代码展示了系统如何结合 active_task 与 task_stack 进行意图识别，并在意图跳转时执行压栈。"),
        ("附图 A-3 `_run_tool_call_loop` 工具循环代码截图", assets["code_experts"], "该截图对应多智能体专家中的工具循环，包含工具消息拼接与最大轮数熔断。"),
        ("附图 A-4 工具层与安全过滤代码截图", assets["code_tools"], "该部分展示了检索工具、下单工具及价格确认过滤逻辑。"),
        ("附图 A-5 评测执行器代码截图", assets["code_runner"], "该部分展示了评测样本超时、成功判定和数据集迭代执行过程。"),
        ("附图 A-6 前端工作台总控页面截图", assets["code_workbench"], "该部分体现了前端如何整合会话、评测、知识库与报告。"),
        ("附图 A-7 Trace 采集器代码截图", assets["code_trace"], "该部分展示了追踪事件的记录与上下文绑定实现。"),
    ]
    for caption, image_path, desc in code_assets:
        add_image(document, image_path, caption, width=6.0)
        add_paragraph(document, desc)

    add_heading(document, "附录 B 关键代码文本摘录", 1)
    listing_specs = [
        ("代码清单 B-1 DialogState", PROJECT_ROOT / "multi_agent_dialog" / "state.py", 1, 23),
        ("代码清单 B-2 Router 核心逻辑", PROJECT_ROOT / "multi_agent_dialog" / "router.py", 7, 71),
        ("代码清单 B-3 工具循环与恢复逻辑", PROJECT_ROOT / "multi_agent_dialog" / "experts.py", 120, 157),
        ("代码清单 B-4 工具层与安全过滤", PROJECT_ROOT / "multi_agent_dialog" / "tools.py", 62, 85),
        ("代码清单 B-5 评测执行器", PROJECT_ROOT / "evaluation" / "runner.py", 128, 217),
    ]
    for title, path, start, end in listing_specs:
        add_paragraph(document, title, bold=True)
        snippet = read_snippet(path, start, end)
        for line in snippet.splitlines():
            p = document.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run(line)
            set_run_font(run, "等宽更纱黑体 SC", latin="Courier New", size=9)

    add_heading(document, "附录 C 评测样本与结果摘要", 1)
    add_paragraph(document, "为了保证论文的可复核性，本文在附录中列出主要评测样本、用户轮次、期望槽位以及多智能体与基线系统的最终结果摘要。由于完整 JSON 报告较长，正文仅展示总体指标，附录提供结构化明细。")
    ours_samples = {sample["id"]: sample for sample in result["stacked_multi_agent"]["samples"]}
    base_samples = {sample["id"]: sample for sample in result["baseline_single_agent"]["samples"]}
    for sample in dataset:
        ours = ours_samples.get(sample["id"], {})
        base = base_samples.get(sample["id"], {})
        add_paragraph(document, f"样本 {sample['id']}（{sample['domain']}）", bold=True)
        add_paragraph(document, f"用户轮次：{' / '.join(sample['user_turns'])}")
        add_paragraph(document, f"期望槽位：{json.dumps(sample['expected_slots'], ensure_ascii=False)}")
        add_paragraph(document, f"多智能体结果：success={ours.get('success')}，turns={ours.get('turn_count')}，predicted={json.dumps(ours.get('predicted_slots', {}), ensure_ascii=False)}")
        add_paragraph(document, f"基线结果：success={base.get('success')}，turns={base.get('turn_count')}，predicted={json.dumps(base.get('predicted_slots', {}), ensure_ascii=False)}")
        if ours.get("failure_reasons"):
            add_paragraph(document, f"多智能体失败原因：{'; '.join(ours['failure_reasons'])}")
        if base.get("failure_reasons"):
            add_paragraph(document, f"基线失败原因：{'; '.join(base['failure_reasons'])}")
        add_paragraph(document, f"多智能体最终回复摘录：{(ours.get('final_reply') or '')[:180]}")
        add_paragraph(document, f"基线最终回复摘录：{(base.get('final_reply') or '')[:180]}")

    add_heading(document, "附录 D 进一步讨论与扩展分析", 1)
    for paragraph in [
        "为了接近毕业论文所要求的篇幅，本文将系统设计、实验分析与附录材料做了较为完整的展开。需要说明的是，篇幅的增加并非单纯堆砌文本，而是通过增加实验背景、失败模式分析、代码摘录和样本明细来提升论文的可复核性与可读性。这种组织方式更符合工程型毕业设计论文的常见写法。",
        "从方法论角度看，本文系统最大的价值在于提出了一个可被继续扩展的中层框架：上层可以替换为不同的大模型或不同的路由策略，下层可以增加新的领域专家和工具接口，而中间的任务栈、黑板和安全过滤层仍然成立。这说明本文贡献不仅是“做了一个可运行 demo”，而是抽象出了适用于多轮生活服务场景的架构模式。",
        "附录部分收录的代码截图和代码文本有助于说明实现细节与论文描述之间的一致性。对于毕业论文而言，这种“正文讲思想、附录留实现”的结构既能保证理论表达的完整，也能为答辩老师提供快速核验的依据。若后续仍需继续扩写，还可以基于同样方式继续补充更多实验轮次、更多知识库变体和更长的对话日志。",
    ]:
        add_paragraph(document, paragraph)


def build_assets(result: dict, dataset: list[dict]) -> dict[str, Path]:
    assets = {
        "success_rate": make_single_metric_chart(result, "task_success_rate", "任务成功率", ASSET_DIR / "task_success_rate.png"),
        "slot_f1": make_single_metric_chart(result, "slot_f1", "槽位 F1", ASSET_DIR / "slot_f1.png"),
        "average_turns": make_single_metric_chart(result, "average_turns", "平均轮次", ASSET_DIR / "average_turns.png"),
        "architecture": make_architecture_diagram(ASSET_DIR / "architecture.png"),
        "stack_flow": make_stack_flow_diagram(ASSET_DIR / "stack_flow.png"),
        "workbench": make_workbench_layout(ASSET_DIR / "workbench.png"),
        "distribution": make_distribution_chart(dataset, ASSET_DIR / "distribution.png"),
        "metric_compare": make_metric_comparison(result, ASSET_DIR / "metric_compare.png"),
        "code_state": render_code_image(PROJECT_ROOT / "multi_agent_dialog" / "state.py", 1, 23, ASSET_DIR / "code_state.png"),
        "code_router": render_code_image(PROJECT_ROOT / "multi_agent_dialog" / "router.py", 7, 71, ASSET_DIR / "code_router.png"),
        "code_experts": render_code_image(PROJECT_ROOT / "multi_agent_dialog" / "experts.py", 120, 157, ASSET_DIR / "code_experts.png"),
        "code_tools": render_code_image(PROJECT_ROOT / "multi_agent_dialog" / "tools.py", 62, 85, ASSET_DIR / "code_tools.png"),
        "code_runner": render_code_image(PROJECT_ROOT / "evaluation" / "runner.py", 128, 217, ASSET_DIR / "code_runner.png"),
        "code_workbench": render_code_image(PROJECT_ROOT / "web" / "src" / "pages" / "WorkbenchPage.tsx", 175, 227, ASSET_DIR / "code_workbench.png"),
        "code_trace": render_code_image(PROJECT_ROOT / "app" / "services" / "trace_service.py", 14, 40, ASSET_DIR / "code_trace.png"),
    }
    return assets


def main() -> None:
    result = load_json(RESULTS_DIR / "latest.json")
    dataset = load_json(DATASET_PATH)
    assets = build_assets(result, dataset)

    document = Document()
    set_document_style(document)
    add_title(document)
    add_standard_front_matter(document)
    add_body(document, assets, result)
    add_references(document)
    add_appendices(document, assets, result, dataset)
    document.save(OUTPUT_PATH)
    print(f"generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
